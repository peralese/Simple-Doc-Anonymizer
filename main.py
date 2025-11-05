import csv
import os
import json
import re
from docx import Document

def load_substitution_csv(file_path):
    substitutions = []
    try:
        with open(file_path, mode='r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            if 'original' not in reader.fieldnames or 'replacement' not in reader.fieldnames:
                print("Error: CSV must contain 'original' and 'replacement' headers.")
                return []
            for row in reader:
                substitutions.append((row['original'].strip(), row['replacement'].strip()))
    except Exception as e:
        print(f"Error reading substitution file: {e}")
        return []
    return substitutions

def get_substitution_map():
    substitutions = []
    print("\nEnter substitution pairs (original -> replacement).")
    while True:
        original = input("Enter original term (or leave blank to finish): ").strip()
        if not original:
            break
        replacement = input(f"Enter replacement for '{original}': ").strip()
        substitutions.append((original, replacement))
    return substitutions

def build_substitution_patterns(substitutions, mode):
    flags = 0
    if mode in ("literal_ci", "regex_ci"):
        flags |= re.IGNORECASE

    compiled = []
    for original, replacement in substitutions:
        if mode in ("literal_cs", "literal_ci"):
            pattern = re.compile(re.escape(original), flags)
        else:  # regex
            try:
                pattern = re.compile(original, flags)
            except re.error as e:
                print(f"Invalid regex '{original}': {e}. Skipping.")
                continue
        compiled.append((pattern, replacement, original))
    return compiled

def process_text(text, compiled_substitutions, stats):
    for pattern, replacement, original in compiled_substitutions:
        text, count = pattern.subn(replacement, text)
        if count:
            stats[original]['count'] += count
    return text

def write_substitution_log(dir_name, name, stats):
    csv_file_path = os.path.join(dir_name, f"{name}_substitution_log.csv")
    with open(csv_file_path, mode='w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['original_term', 'replacement_term', 'occurrences'])
        for original in stats:
            writer.writerow([original, stats[original]['replacement'], stats[original]['count']])
    return csv_file_path

def process_docx(file_path, compiled_substitutions, stats):
    print("\nLoading DOCX document...")
    doc = Document(file_path)

    for para in doc.paragraphs:
        para.text = process_text(para.text, compiled_substitutions, stats)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = process_text(cell.text, compiled_substitutions, stats)

    for section in doc.sections:
        for para in section.header.paragraphs:
            para.text = process_text(para.text, compiled_substitutions, stats)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = process_text(cell.text, compiled_substitutions, stats)
        for para in section.footer.paragraphs:
            para.text = process_text(para.text, compiled_substitutions, stats)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = process_text(cell.text, compiled_substitutions, stats)

    dir_name, base_name = os.path.split(file_path)
    name, ext = os.path.splitext(base_name)
    new_file_path = os.path.join(dir_name, f"{name}_anonymized{ext}")
    doc.save(new_file_path)
    return new_file_path

def process_txt(file_path, compiled_substitutions, stats):
    print("\nLoading TXT file...")
    with open(file_path, mode='r', encoding='utf-8') as f:
        content = f.read()
    content = process_text(content, compiled_substitutions, stats)
    dir_name, base_name = os.path.split(file_path)
    name, ext = os.path.splitext(base_name)
    new_file_path = os.path.join(dir_name, f"{name}_anonymized{ext}")
    with open(new_file_path, mode='w', encoding='utf-8') as f:
        f.write(content)
    return new_file_path

def _process_json_values(node, compiled_substitutions, stats, anonymize_keys=False):
    if isinstance(node, dict):
        new_obj = {}
        for k, v in node.items():
            new_key = process_text(k, compiled_substitutions, stats) if anonymize_keys and isinstance(k, str) else k
            new_obj[new_key] = _process_json_values(v, compiled_substitutions, stats, anonymize_keys)
        return new_obj
    elif isinstance(node, list):
        return [_process_json_values(item, compiled_substitutions, stats, anonymize_keys) for item in node]
    elif isinstance(node, str):
        return process_text(node, compiled_substitutions, stats)
    else:
        return node

def process_json(file_path, compiled_substitutions, stats):
    print("\nLoading JSON file...")
    try:
        with open(file_path, mode='r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        print(f"Error reading JSON file: {e}")
        return None

    anonymize_keys = input("Also anonymize string keys in JSON? (y/n): ").strip().lower() == 'y'
    data = _process_json_values(data, compiled_substitutions, stats, anonymize_keys=anonymize_keys)

    dir_name, base_name = os.path.split(file_path)
    name, ext = os.path.splitext(base_name)
    new_file_path = os.path.join(dir_name, f"{name}_anonymized{ext}")
    try:
        with open(new_file_path, mode='w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Error writing JSON file: {e}")
        return None
    return new_file_path

def anonymize_file():
    print("=== File Anonymizer (.docx, .txt, .json) ===")

    # 1. Get file path
    file_path = input("Enter the full path to the file: ").strip()
    if not os.path.isfile(file_path):
        print("Error: File not found.")
        return

    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    if ext not in {'.docx', '.txt', '.json'}:
        print("Error: Unsupported file type. Supported: .docx, .txt, .json")
        return

    # 2. Get substitutions
    use_csv = input("Do you have a CSV file with substitutions? (y/n): ").strip().lower()
    if use_csv == 'y':
        csv_path = input("Enter the path to the CSV file: ").strip()
        substitutions = load_substitution_csv(csv_path)
        if not substitutions:
            print("No valid substitutions loaded from file. Exiting.")
            return
    else:
        substitutions = get_substitution_map()
        if not substitutions:
            print("No substitutions provided. Exiting.")
            return

    # 3. Choose matching mode
    print("\nChoose matching mode:")
    print("  1) literal (case-sensitive)")
    print("  2) literal (case-insensitive)")
    print("  3) regex (case-sensitive)")
    print("  4) regex (case-insensitive)")
    choice = input("Enter 1, 2, 3, or 4: ").strip()
    mode_map = {'1': 'literal_cs', '2': 'literal_ci', '3': 'regex_cs', '4': 'regex_ci'}
    mode = mode_map.get(choice, 'literal_cs')

    # 4. Initialize tracking and compile patterns
    stats = {orig: {'replacement': repl, 'count': 0} for orig, repl in substitutions}
    compiled = build_substitution_patterns(substitutions, mode)
    if not compiled:
        print("No valid substitutions after compiling patterns. Exiting.")
        return

    # 5. Process based on extension
    if ext == '.docx':
        new_file_path = process_docx(file_path, compiled, stats)
    elif ext == '.txt':
        new_file_path = process_txt(file_path, compiled, stats)
    elif ext == '.json':
        new_file_path = process_json(file_path, compiled, stats)
    else:
        print("Unexpected error: unsupported extension after validation.")
        return

    if not new_file_path:
        print("Anonymization failed.")
        return

    # 6. Save substitution log as CSV
    dir_name, base_name = os.path.split(file_path)
    name, _ = os.path.splitext(base_name)
    csv_file_path = write_substitution_log(dir_name, name, stats)

    # 7. Done
    print("\n=== Anonymization Complete ===")
    print(f"Anonymized file saved to: {new_file_path}")
    print(f"Substitution log saved to: {csv_file_path}")
    print("\nSummary of replacements:")
    for original in stats:
        print(f"  '{original}' -> '{stats[original]['replacement']}': {stats[original]['count']} occurrence(s)")

if __name__ == "__main__":
    anonymize_file()

